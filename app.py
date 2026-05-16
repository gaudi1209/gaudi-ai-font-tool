# -*- coding: utf-8 -*-
"""AI字体生产工具 - Flask主应用"""

import os
import sys
import subprocess

sys.stdout.reconfigure(encoding='utf-8')

from flask import Flask, render_template, request, jsonify, send_from_directory
from config import HOST, PORT, DEBUG, ZI2ZI_DIR, DEFAULT_BASE_CHECKPOINT, DEFAULT_SOURCE_FONT, DEFAULT_REF_FONT, TRAIN_DEFAULTS, GENERATE_DEFAULTS
from utils.train_manager import train_manager
from utils.generate_manager import generate_manager
from utils.ocr_manager import ocr_manager

app = Flask(__name__)

# ===== 页面路由 =====

@app.route('/')
def index():
    return render_template('train.html')

@app.route('/train')
def train_page():
    return render_template('train.html')

@app.route('/generate')
def generate_page():
    return render_template('generate.html')

@app.route('/ocr')
def ocr_page():
    return render_template('ocr.html')


# ===== 训练API =====

@app.route('/api/train/prepare', methods=['POST'])
def train_prepare():
    """准备训练数据"""
    params = request.json
    ref_font = params.get('ref_font', '')
    source_font = params.get('source_font', '')
    output_dir = params.get('output_dir', '')
    char_count = params.get('char_count')  # None=全部, 数字=限制数量

    if not ref_font or not source_font or not output_dir:
        return jsonify({"success": False, "error": "请填写完整路径"})

    # 自动设置data_path和test_npz_path
    result = train_manager.prepare_data(output_dir, ref_font, source_font, char_count=char_count)
    return jsonify(result)


@app.route('/api/train/start', methods=['POST'])
def train_start():
    """启动训练"""
    params = request.json

    # 补充默认值
    params.setdefault('base_checkpoint', DEFAULT_BASE_CHECKPOINT)
    params.setdefault('source_font', DEFAULT_SOURCE_FONT)
    params.setdefault('model', TRAIN_DEFAULTS['model'])
    params.setdefault('sampling_method', TRAIN_DEFAULTS['sampling_method'])
    params.setdefault('num_sampling_steps', TRAIN_DEFAULTS['num_sampling_steps'])
    params.setdefault('lora_targets', TRAIN_DEFAULTS['lora_targets'])

    # 自动推断data_path和test_npz
    output_dir = params.get('output_dir', '')
    if output_dir and 'data_path' not in params:
        params['data_path'] = output_dir
    if output_dir and 'test_npz_path' not in params:
        test_npz = os.path.join(output_dir, 'test.npz')
        font_dir = os.path.join(output_dir, '001_font')
        if not os.path.isdir(font_dir):
            return jsonify({"success": False, "error": f"数据目录不存在: {font_dir}，请先准备数据"})
        if not os.path.exists(test_npz):
            return jsonify({"success": False, "error": f"test.npz不存在，请先准备数据"})
        params['test_npz_path'] = test_npz

    result = train_manager.start_training(params)
    return jsonify(result)


@app.route('/api/train/test_generate', methods=['POST'])
def train_test_generate():
    """训练完成后快速测试生成"""
    params = request.json
    output_dir = params.get('output_dir', '')
    ref_font = params.get('ref_font', '')
    source_font = params.get('source_font', '')

    if not output_dir or not ref_font or not source_font:
        return jsonify({"success": False, "error": "请填写输出目录、学习字库和源字体路径"})

    result = train_manager.test_generate(output_dir, ref_font, source_font)
    return jsonify(result)


@app.route('/api/train/stop', methods=['POST'])
def train_stop():
    return jsonify(train_manager.stop_training())


@app.route('/api/train/status')
def train_status():
    return jsonify(train_manager.get_status())


@app.route('/api/train/logs')
def train_logs():
    offset = request.args.get('offset', 0, type=int)
    limit = request.args.get('limit', 100, type=int)
    return jsonify(train_manager.get_logs(offset, limit))


# ===== 生成API =====

@app.route('/api/generate/start', methods=['POST'])
def generate_start():
    """启动迭代生成"""
    params = request.json
    params.setdefault('source_font', DEFAULT_SOURCE_FONT)
    params.setdefault('ref_font', DEFAULT_REF_FONT)
    params.setdefault('resolution', GENERATE_DEFAULTS['resolution'])
    params.setdefault('ref_size', GENERATE_DEFAULTS['ref_size'])
    params.setdefault('batch_size', GENERATE_DEFAULTS['batch_size'])

    result = generate_manager.start_generate(params)
    return jsonify(result)


@app.route('/api/generate/stop', methods=['POST'])
def generate_stop():
    return jsonify(generate_manager.stop_generate())


@app.route('/api/generate/status')
def generate_status():
    return jsonify(generate_manager.get_status())


def _check_pua_system_fonts(pua_codes):
    """检查系统字体对PUA字符的支持，尝试从字形名推断标准Unicode映射"""
    import re
    from fontTools.ttLib import TTFont

    SYSTEM_FONTS = [
        ('宋体', 'C:\\Windows\\Fonts\\simsun.ttc', 0),
        ('楷体', 'C:\\Windows\\Fonts\\simkai.ttf', None),
        ('仿宋', 'C:\\Windows\\Fonts\\simfang.ttf', None),
        ('微软雅黑', 'C:\\Windows\\Fonts\\msyh.ttc', 0),
        ('明體', 'C:\\Windows\\Fonts\\mingliu.ttc', 0),
    ]

    results = {}
    for name, path, fnum in SYSTEM_FONTS:
        if not os.path.isfile(path):
            continue
        try:
            kw = {'fontNumber': fnum} if fnum is not None else {}
            font = TTFont(path, **kw)
            cmap = font.getBestCmap()
            font.close()
        except Exception:
            continue

        for cp in pua_codes:
            if cp not in cmap:
                continue
            glyph_name = cmap[cp]
            if cp not in results:
                results[cp] = {'fonts': [], 'mapped': None, 'mapped_code': None}
            results[cp]['fonts'].append(name)

            # 尝试从字形名推断标准Unicode (如 uni7D2C → U+7D2C 紬)
            if results[cp]['mapped'] is None:
                m = re.match(r'^uni([0-9A-Fa-f]{4,6})$', glyph_name, re.IGNORECASE)
                if not m:
                    m = re.match(r'^u([0-9A-Fa-f]{5,6})$', glyph_name, re.IGNORECASE)
                if m:
                    mapped = int(m.group(1), 16)
                    if mapped != cp and not (0xE000 <= mapped <= 0xF8FF):
                        try:
                            results[cp]['mapped'] = chr(mapped)
                            results[cp]['mapped_code'] = f'U+{mapped:04X}'
                        except Exception:
                            pass
    return results


@app.route('/api/generate/missing_chars', methods=['POST'])
def missing_chars():
    """计算TTF相对文本文件或文本内容的缺失字符"""
    params = request.json
    ttf_path = params.get('ttf_path', '')
    text = params.get('text', '')
    text_file = params.get('text_file', '')

    # 优先读取文本文件
    if text_file and os.path.isfile(text_file):
        try:
            ext = os.path.splitext(text_file)[1].lower()
            if ext == '.txt':
                with open(text_file, 'r', encoding='utf-8') as f:
                    text = f.read()
            elif ext == '.docx':
                from docx import Document
                doc = Document(text_file)
                text = '\n'.join(p.text for p in doc.paragraphs)
            elif ext == '.pdf':
                import fitz  # PyMuPDF
                doc = fitz.open(text_file)
                text = ''.join(page.get_text() for page in doc)
                doc.close()
            elif ext == '.epub':
                import zipfile
                from xml.etree import ElementTree
                with zipfile.ZipFile(text_file) as zf:
                    for name in zf.namelist():
                        if name.endswith(('.html', '.xhtml', '.htm')):
                            with zf.open(name) as f:
                                tree = ElementTree.parse(f)
                                root = tree.getroot()
                                for elem in root.iter():
                                    if elem.text:
                                        text += elem.text + ' '
            elif ext == '.mobi':
                import subprocess as sp
                result = sp.run(['python', '-m', 'ebooklib.mobi', text_file],
                                capture_output=True, text=True, encoding='utf-8')
                if result.returncode != 0:
                    # fallback: 尝试作为 zip 处理
                    try:
                        import zipfile
                        with zipfile.ZipFile(text_file) as zf:
                            for name in zf.namelist():
                                if name.endswith(('.html', '.htm')):
                                    with zf.open(name) as f:
                                        text += f.read().decode('utf-8', errors='ignore')
                    except Exception:
                        return jsonify({"success": False, "error": f"不支持 MOBI 格式，建议先用 Calibre 转为 TXT: {result.stderr}"})
            else:
                return jsonify({"success": False, "error": f"不支持的文件格式: {ext}"})
        except Exception as e:
            return jsonify({"success": False, "error": f"读取文件失败: {e}"})

    if not ttf_path or not text:
        return jsonify({"success": False, "error": "请填写TTF路径和文本文件"})

    try:
        from utils.charset_utils import get_font_chars, is_cjk
        font_chars = get_font_chars(ttf_path)
        text_chars = set(ord(c) for c in text if is_cjk(ord(c)))
        missing = sorted(text_chars - font_chars)

        # PUA 字符检测 (U+E000-U+F8FF)
        pua_map = {}
        for c in text:
            cp = ord(c)
            if 0xE000 <= cp <= 0xF8FF:
                pua_map[cp] = pua_map.get(cp, 0) + 1
        pua_sorted = sorted(pua_map.keys())

        # 系统字体检查
        pua_sys = _check_pua_system_fonts(pua_sorted) if pua_sorted else {}
        pua_list = []
        for c in pua_sorted:
            info = pua_sys.get(c, {'fonts': [], 'mapped': None, 'mapped_code': None})
            pua_list.append({
                "code": f"U+{c:04X}",
                "char": chr(c),
                "count": pua_map[c],
                "system_fonts": info['fonts'],
                "mapped_char": info['mapped'],
                "mapped_code": info['mapped_code'],
            })

        return jsonify({
            "success": True,
            "missing_chars": [chr(c) for c in missing],
            "missing_count": len(missing),
            "text_total": len(text_chars),
            "font_total": len(font_chars),
            "pua_chars": pua_list,
            "pua_count": len(pua_sorted),
            "pua_total": sum(pua_map.values()),
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


@app.route('/api/generate/diff_groups', methods=['POST'])
def diff_groups():
    """计算TTF与字符集的差集分组，输出分组txt到输出目录"""
    params = request.json
    ttf_path = params.get('ttf_path', '')
    charset_name = params.get('charset', 'GB2312')
    output_dir = params.get('output_dir', '')
    export_group_size = params.get('export_group_size', 0)  # 导出txt的分组字数，0=不导出

    if not ttf_path:
        return jsonify({"success": False, "error": "请填写TTF路径"})

    try:
        from utils.charset_utils import get_font_chars, get_charset, split_into_groups
        font_chars = get_font_chars(ttf_path)
        charset = get_charset(charset_name)
        missing = sorted(charset - font_chars)
        groups = split_into_groups(missing, 500)

        # 导出带分组的txt文件（所有分组写在一个文件中）
        export_file = None
        if export_group_size > 0 and output_dir and os.path.isdir(output_dir):
            from datetime import datetime
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            txt_name = f"missing_{charset_name}_{timestamp}.txt"
            txt_path = os.path.join(output_dir, txt_name)
            export_groups = split_into_groups(missing, export_group_size)
            with open(txt_path, 'w', encoding='utf-8') as f:
                for i, g in enumerate(export_groups):
                    group_label = f"第{i + 1}组 ({len(g)}字)"
                    f.write(f"{'=' * 40}\n")
                    f.write(f"{group_label}\n")
                    f.write(f"{'=' * 40}\n")
                    f.write(''.join(chr(c) for c in g))
                    f.write('\n\n')
            export_file = txt_name

        return jsonify({
            "success": True,
            "charset_name": charset_name,
            "charset_size": len(charset),
            "font_size": len(font_chars & charset),
            "missing_count": len(missing),
            "groups": [{"index": i, "size": len(g), "chars": [chr(c) for c in g]} for i, g in enumerate(groups)],
            "export_file": export_file,
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


@app.route('/api/generate/pua_check', methods=['POST'])
def pua_check():
    """检测文本中的 PUA 字符 (U+E000-U+F8FF)"""
    params = request.json
    text = params.get('text', '')
    text_file = params.get('text_file', '')

    # 优先读取文本文件
    if text_file and os.path.isfile(text_file):
        try:
            ext = os.path.splitext(text_file)[1].lower()
            if ext == '.txt':
                with open(text_file, 'r', encoding='utf-8') as f:
                    text = f.read()
            elif ext == '.docx':
                from docx import Document
                doc = Document(text_file)
                text = '\n'.join(p.text for p in doc.paragraphs)
            elif ext == '.pdf':
                import fitz
                doc = fitz.open(text_file)
                text = ''.join(page.get_text() for page in doc)
                doc.close()
            elif ext == '.epub':
                import zipfile
                from xml.etree import ElementTree
                with zipfile.ZipFile(text_file) as zf:
                    for name in zf.namelist():
                        if name.endswith(('.html', '.xhtml', '.htm')):
                            with zf.open(name) as f:
                                tree = ElementTree.parse(f)
                                root = tree.getroot()
                                for elem in root.iter():
                                    if elem.text:
                                        text += elem.text + ' '
            else:
                return jsonify({"success": False, "error": f"不支持的文件格式: {ext}"})
        except Exception as e:
            return jsonify({"success": False, "error": f"读取文件失败: {e}"})

    if not text:
        return jsonify({"success": False, "error": "请提供文本或文本文件"})

    try:
        pua_set = {}
        for c in text:
            cp = ord(c)
            if 0xE000 <= cp <= 0xF8FF:
                pua_set[cp] = pua_set.get(cp, 0) + 1
        pua_sorted = sorted(pua_set.keys())

        # 系统字体检查
        pua_sys = _check_pua_system_fonts(pua_sorted) if pua_sorted else {}
        pua_list = []
        for c in pua_sorted:
            info = pua_sys.get(c, {'fonts': [], 'mapped': None, 'mapped_code': None})
            pua_list.append({
                "code": f"U+{c:04X}",
                "char": chr(c),
                "count": pua_set[c],
                "system_fonts": info['fonts'],
                "mapped_char": info['mapped'],
                "mapped_code": info['mapped_code'],
            })

        return jsonify({
            "success": True,
            "pua_chars": pua_list,
            "pua_count": len(pua_sorted),
            "pua_total": sum(pua_set.values()),
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


@app.route('/api/generate/images')
def generate_images():
    """获取生成结果图片列表"""
    import glob
    output_dir = request.args.get('dir', '')
    if not output_dir or not os.path.isdir(output_dir):
        return jsonify({"images": []})

    images = []
    from urllib.parse import quote
    for f in sorted(glob.glob(os.path.join(output_dir, '*.png')))[:200]:
        fname = os.path.basename(f)
        label = ''
        if '_' in fname:
            label = fname.split('_')[1].replace('.png', '')

        images.append({
            "name": fname,
            "url": f"/api/image?path={quote(f)}",
            "label": label,
        })

    return jsonify({"images": images})


@app.route('/api/generate/add_suffix', methods=['POST'])
def add_hanzi_suffix():
    """给目录中的图片添加汉字后缀：uni4E00.png → uni4E00_一.png"""
    import re
    directory = request.json.get('dir', '')
    if not directory or not os.path.isdir(directory):
        return jsonify({"success": False, "error": "目录不存在"})

    renamed = 0
    skipped = 0
    for f in sorted(os.listdir(directory)):
        if not f.endswith('.png'):
            continue
        if re.search(r'^[uU][nN]?[0-9A-Fa-f]+_.+\.png$', f):
            skipped += 1
            continue
        m = re.match(r'^(uni|u)([0-9A-Fa-f]+)\.png$', f)
        if not m:
            continue
        prefix = m.group(1)
        code_hex = m.group(2)
        code = int(code_hex, 16)
        try:
            char = chr(code)
            new_name = f"{prefix}{code_hex}_{char}.png"
            src = os.path.join(directory, f)
            dst = os.path.join(directory, new_name)
            if not os.path.exists(dst):
                os.rename(src, dst)
                renamed += 1
            else:
                skipped += 1
        except (ValueError, OverflowError):
            skipped += 1

    return jsonify({"success": True, "renamed": renamed, "skipped": skipped})


@app.route('/api/generate/remove_suffix', methods=['POST'])
def remove_hanzi_suffix():
    """去掉目录中图片的汉字后缀：uni4E00_一.png → uni4E00.png"""
    import re
    directory = request.json.get('dir', '')
    if not directory or not os.path.isdir(directory):
        return jsonify({"success": False, "error": "目录不存在"})

    renamed = 0
    skipped = 0
    for f in sorted(os.listdir(directory)):
        if not f.endswith('.png'):
            continue
        m = re.match(r'^(uni|u)([0-9A-Fa-f]+)_.+\.png$', f)
        if not m:
            skipped += 1
            continue
        prefix = m.group(1)
        code_hex = m.group(2)
        new_name = f"{prefix}{code_hex}.png"
        src = os.path.join(directory, f)
        dst = os.path.join(directory, new_name)
        if not os.path.exists(dst):
            os.rename(src, dst)
            renamed += 1
        else:
            skipped += 1

    return jsonify({"success": True, "renamed": renamed, "skipped": skipped})


@app.route('/api/image')
def serve_image():
    """提供图片文件，支持后缀名自动回退"""
    import re
    filepath = request.args.get('path', '')
    if not filepath:
        return "Not found", 404

    # 直接找到就返回
    if os.path.isfile(filepath):
        return send_from_directory(os.path.dirname(filepath), os.path.basename(filepath))

    # 找不到时尝试去掉汉字后缀：uni9B74_魴.png → uni9B74.png
    basename = os.path.basename(filepath)
    m = re.match(r'^(uni|u)([0-9A-Fa-f]+)_.+\.png$', basename, re.IGNORECASE)
    if m:
        alt_name = f"{m.group(1)}{m.group(2)}.png"
        alt_path = os.path.join(os.path.dirname(filepath), alt_name)
        if os.path.isfile(alt_path):
            return send_from_directory(os.path.dirname(alt_path), alt_name)

    # 再尝试加汉字后缀：uni9B74.png → 查找同目录下 uni9B74_*.png
    m2 = re.match(r'^(uni|u)([0-9A-Fa-f]+)\.png$', basename, re.IGNORECASE)
    if m2:
        prefix = f"{m2.group(1)}{m2.group(2)}_"
        parent = os.path.dirname(filepath)
        try:
            for f in os.listdir(parent):
                if f.startswith(prefix) and f.endswith('.png'):
                    return send_from_directory(parent, f)
        except OSError:
            pass

    return "Not found", 404


@app.route('/api/delete_file', methods=['POST'])
def delete_file():
    """删除指定文件，支持后缀名自动回退"""
    import re
    filepath = request.json.get('path', '')
    if not filepath:
        return jsonify({"success": False, "error": "未指定文件"})

    # 直接找到就删除
    if os.path.isfile(filepath):
        os.unlink(filepath)
        return jsonify({"success": True})

    # 尝试去掉汉字后缀：uni9B74_魴.png → uni9B74.png
    basename = os.path.basename(filepath)
    m = re.match(r'^(uni|u)([0-9A-Fa-f]+)_.+\.png$', basename, re.IGNORECASE)
    if m:
        alt_name = f"{m.group(1)}{m.group(2)}.png"
        alt_path = os.path.join(os.path.dirname(filepath), alt_name)
        if os.path.isfile(alt_path):
            os.unlink(alt_path)
            return jsonify({"success": True})

    # 尝试查找带后缀版本：uni9B74.png → uni9B74_*.png
    m2 = re.match(r'^(uni|u)([0-9A-Fa-f]+)\.png$', basename, re.IGNORECASE)
    if m2:
        prefix = f"{m2.group(1)}{m2.group(2)}_"
        parent = os.path.dirname(filepath)
        try:
            for f in os.listdir(parent):
                if f.startswith(prefix) and f.endswith('.png'):
                    os.unlink(os.path.join(parent, f))
                    return jsonify({"success": True})
        except OSError:
            pass

    return jsonify({"success": False, "error": "文件不存在"})


# ===== OCR API =====

@app.route('/api/ocr/start', methods=['POST'])
def ocr_start():
    params = request.json
    result = ocr_manager.start_ocr(params)
    return jsonify(result)


@app.route('/api/ocr/stop', methods=['POST'])
def ocr_stop():
    return jsonify(ocr_manager.stop_ocr())


@app.route('/api/ocr/status')
def ocr_status():
    return jsonify(ocr_manager.get_status())


@app.route('/api/ocr/results')
def ocr_results():
    limit = request.args.get('limit', 100, type=int)
    offset = request.args.get('offset', 0, type=int)
    return jsonify(ocr_manager.get_results(offset, limit))


# ===== 通用API =====

@app.route('/api/open_dir', methods=['POST'])
def open_dir():
    """打开目录或文件"""
    path = request.json.get('path', '')
    if not path:
        return jsonify({"success": False, "error": "未指定路径"})
    try:
        import subprocess, ctypes, time
        subprocess.Popen(f'explorer "{path}"', shell=True)
        # 等待窗口出现后置前
        time.sleep(0.5)
        ctypes.windll.user32.keybd_event(0x12, 0, 0, 0)  # Alt press
        ctypes.windll.user32.keybd_event(0x12, 0, 2, 0)  # Alt release
        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


@app.route('/api/browse', methods=['POST'])
def browse_path():
    """浏览文件/目录 - 返回候选列表"""
    print(f"[browse] request: {request.json}")
    path = request.json.get('path', '')
    file_type = request.json.get('type', 'dir')  # dir / file
    extensions = request.json.get('extensions', [])

    # 如果路径是文件直接返回
    if path and os.path.isfile(path):
        return jsonify({"success": True, "selected": path, "type": "file"})

    # 目标目录
    target = path if path and os.path.isdir(path) else os.path.dirname(path) if path else ''

    if not target:
        target = os.path.expanduser('~')

    if not os.path.isdir(target):
        return jsonify({"success": False, "error": "路径不存在"})

    items = []

    # 盘符入口（Windows）
    import string
    drives = []
    for letter in string.ascii_uppercase:
        drive = letter + ':\\'
        if os.path.isdir(drive):
            drives.append({"name": drive, "path": drive, "type": "drive"})
    if drives:
        items.append({"name": "📦 盘符", "path": "", "type": "drives", "drives": drives})

    parent = os.path.dirname(target)
    if parent and parent != target:
        items.append({"name": ".. (上级目录)", "path": parent, "type": "parent"})

    try:
        for entry in sorted(os.listdir(target)):
            full = os.path.join(target, entry)
            if os.path.isdir(full):
                items.append({"name": entry, "path": full, "type": "dir"})
            elif file_type == 'file':
                if not extensions or any(entry.lower().endswith(ext) for ext in extensions):
                    items.append({"name": entry, "path": full, "type": "file"})
    except PermissionError:
        pass

    return jsonify({
        "success": True,
        "current": target,
        "items": items,
    })


if __name__ == '__main__':
    print(f"AI字体生产工具 启动于 http://localhost:{PORT}")
    app.run(host=HOST, port=PORT, debug=DEBUG)
